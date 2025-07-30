"""
このファイルは、画面表示以外の様々な関数定義のファイルです。
（Faiss保存方法修正版）
"""

############################################################
# ライブラリの読み込み
############################################################
import os
from dotenv import load_dotenv
import streamlit as st
import logging
import sys
import unicodedata
import gspread
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder, PromptTemplate
from langchain.schema import HumanMessage, AIMessage, Document
from langchain_openai import OpenAIEmbeddings
# ChromaDBの代わりにFaissを使用
from langchain_community.vectorstores import FAISS
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_community.callbacks.streamlit import StreamlitCallbackHandler
from typing import List
from sudachipy import tokenizer, dictionary
from langchain_community.agent_toolkits import SlackToolkit
from langchain.agents import AgentType, initialize_agent
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from docx import Document
from langchain.output_parsers import CommaSeparatedListOutputParser
from langchain import LLMChain
import datetime
import constants as ct
from pathlib import Path
from collections import Counter
from langchain.schema.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter as RegexTextSplitter
import pickle
import hashlib
import json
from oauth2client.service_account import ServiceAccountCredentials
from langchain_community.utilities import GoogleSearchAPIWrapper
from slack_sdk import WebClient
from slack_sdk.errors import SlackApiError
import traceback
from langchain.callbacks.base import BaseCallbackHandler
from typing import Any, Dict, List, Optional, Union


# ============================================================================
# 同義語辞書（必要に応じて拡張可能）
# ============================================================================

SYNONYM_DICT = {
    "受賞": ["受賞歴", "表彰", "アワード", "賞", "栄誉"],
    "実績": ["成果", "業績", "結果", "成績"],
    "会社": ["企業", "法人", "組織", "事業者"],
    "環境": ["エコ", "グリーン", "サステナブル", "持続可能"],
    "製品": ["商品", "サービス", "プロダクト"],
    "顧客": ["お客様", "利用者", "ユーザー", "クライアント"],
    "料金": ["価格", "費用", "コスト", "値段", "金額"],
    "配送": ["発送", "配達", "お届け", "輸送"],
    "注文": ["オーダー", "発注", "購入", "申込"]
}

############################################################
# 設定関連
############################################################
load_dotenv()


############################################################
# Faiss専用の保存・読み込み関数
############################################################

def save_faiss_index(db, base_path):
    """
    Faissインデックスを適切な方法で保存
    
    Args:
        db: FAISSベクトルストアオブジェクト
        base_path: 保存先のベースパス（拡張子なし）
    """
    try:
        # Faiss専用の保存方法を使用
        db.save_local(base_path)
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.info(f"✅ Faissインデックス保存完了: {base_path}")
        return True
    except Exception as e:
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.error(f"❌ Faissインデックス保存エラー: {e}")
        return False

def load_faiss_index(base_path, embeddings):
    """
    Faissインデックスを適切な方法で読み込み
    
    Args:
        base_path: 読み込み元のベースパス（拡張子なし）
        embeddings: 埋め込みオブジェクト
        
    Returns:
        FAISSベクトルストアオブジェクト または None
    """
    try:
        # Faiss専用の読み込み方法を使用
        db = FAISS.load_local(base_path, embeddings, allow_dangerous_deserialization=True)
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.info(f"✅ Faissインデックス読み込み完了: {base_path}")
        return db
    except Exception as e:
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.warning(f"⚠️ Faissインデックス読み込み失敗: {e}")
        return None

def calculate_docs_hash(docs):
    """
    ドキュメントリストのハッシュ値を計算（変更検知用）
    
    Args:
        docs: ドキュメントリスト
        
    Returns:
        ハッシュ値文字列
    """
    # ドキュメントの内容からハッシュを生成
    content_str = ""
    for doc in docs:
        content_str += doc.page_content + str(doc.metadata)
    
    return hashlib.md5(content_str.encode('utf-8')).hexdigest()

def should_rebuild_index(base_path, docs):
    """
    インデックスの再構築が必要かチェック
    
    Args:
        base_path: インデックスのベースパス
        docs: 現在のドキュメントリスト
        
    Returns:
        bool: 再構築が必要な場合True
    """
    metadata_file = f"{base_path}_metadata.json"
    
    # メタデータファイルが存在しない場合は再構築
    if not os.path.exists(metadata_file):
        return True
    
    # Faissインデックスファイルが存在しない場合は再構築
    if not os.path.exists(f"{base_path}.faiss"):
        return True
    
    try:
        # メタデータを読み込み
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        
        # ドキュメントのハッシュを比較
        current_hash = calculate_docs_hash(docs)
        stored_hash = metadata.get('docs_hash', '')
        
        if current_hash != stored_hash:
            logger = logging.getLogger(ct.LOGGER_NAME)
            logger.info("📝 ドキュメントが変更されたため、インデックスを再構築します")
            return True
        
        return False
        
    except Exception as e:
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.warning(f"⚠️ メタデータ読み込みエラー、再構築します: {e}")
        return True

def save_index_metadata(base_path, docs):
    """
    インデックスのメタデータを保存
    
    Args:
        base_path: インデックスのベースパス
        docs: ドキュメントリスト
    """
    try:
        metadata_file = f"{base_path}_metadata.json"
        metadata = {
            'docs_hash': calculate_docs_hash(docs),
            'doc_count': len(docs),
            'created_at': datetime.datetime.now().isoformat(),
            'version': '1.0'
        }
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
            
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.info(f"✅ メタデータ保存完了: {metadata_file}")
        
    except Exception as e:
        logger = logging.getLogger(ct.LOGGER_NAME)
        logger.error(f"❌ メタデータ保存エラー: {e}")

############################################################
# 関数定義
############################################################
def create_base_vectorstore(db_name):
    """
    共通のベクトルストア作成処理
    
    Args:
        db_name: ベクトルDBの保存先ディレクトリ名
        
    Returns:
        FAISSベクトルストアオブジェクト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🔨 ベクトルストア作成開始: {db_name}")
    
    docs_all = []
    
    # AIエージェント機能を使わない場合の処理
    if db_name == ct.DB_ALL_PATH:
        folders = os.listdir(ct.RAG_TOP_FOLDER_PATH)
        # 「data」フォルダ直下の各フォルダ名に対して処理
        for folder_path in folders:
            if folder_path.startswith("."):
                continue
            # フォルダ内の各ファイルのデータをリストに追加
            add_docs(f"{ct.RAG_TOP_FOLDER_PATH}/{folder_path}", docs_all)
    # AIエージェント機能を使う場合の処理
    else:
        # データベース名に対応した、RAG化対象のデータ群が格納されているフォルダパスを取得
        folder_path = ct.DB_NAMES[db_name]
        # フォルダ内の各ファイルのデータをリストに追加
        add_docs(folder_path, docs_all)

    # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
    for doc in docs_all:
        doc.page_content = adjust_string(doc.page_content)
        for key in doc.metadata:
            doc.metadata[key] = adjust_string(doc.metadata[key])
    
    text_splitter = CharacterTextSplitter(
        chunk_size=ct.CHUNK_SIZE,
        chunk_overlap=ct.CHUNK_OVERLAP,
        separator="\n",
    )
    splitted_docs = text_splitter.split_documents(docs_all)
    
    # ✅ チャンク先頭にメタ情報を付加（retrieverと同じ処理）
    for doc in splitted_docs:
        file_name = doc.metadata.get("file_name", "不明")
        category = doc.metadata.get("category", "不明")
        heading = doc.metadata.get("first_heading", "")
        keywords_str = doc.metadata.get("top_keywords", "") 

        prefix = f"【カテゴリ: {category}】【ファイル名: {file_name}】"
        if heading:
            prefix += f"【見出し: {heading}】"
        if keywords_str:  # ← 修正：文字列をそのまま使用
            prefix += f"【キーワード: {keywords_str}】"

        doc.page_content = prefix + "\n" + doc.page_content

    embeddings = OpenAIEmbeddings()

    # Faissインデックスの作成（修正版）
    base_path = f"{db_name}_faiss"
    
    # 再構築が必要かチェック
    if should_rebuild_index(base_path, splitted_docs):
        logger.info(f"🔨 Faissインデックスを構築中: {base_path}")
        db = FAISS.from_documents(splitted_docs, embeddings)
        
        # Faiss専用の保存方法を使用
        if save_faiss_index(db, base_path):
            save_index_metadata(base_path, splitted_docs)
        
    else:
        # 既存のインデックスを読み込み
        db = load_faiss_index(base_path, embeddings)
        if db is None:
            # 読み込み失敗時は再作成
            logger.info("🔄 インデックス読み込み失敗、再作成します")
            db = FAISS.from_documents(splitted_docs, embeddings)
            if save_faiss_index(db, base_path):
                save_index_metadata(base_path, splitted_docs)
    
    logger.info(f"✅ ベクトルストア作成完了: {db_name}")
    return db

def run_doc_chain_base(chain_name, param):
    """
    共通のDocチェーン実行処理
    
    Args:
        chain_name: チェーン名（例: "company", "service", "customer" など）
        param: ユーザー入力値
        
    Returns:
        LLMからの回答
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🔗 {chain_name}チェーン実行開始")
    
    try:
        # チェーン名から対応するsession_stateの属性を取得
        # 例: chain_name="company" → st.session_state.company_doc_chain
        chain_attr_name = f"{chain_name}_doc_chain"
        chain = getattr(st.session_state, chain_attr_name)
        
        # チェーン実行
        ai_msg = chain.invoke({
            "input": param,
            "chat_history": st.session_state.chat_history
        })
        
        # 会話履歴への追加
        st.session_state.chat_history.extend([
            HumanMessage(content=param),
            AIMessage(content=ai_msg["answer"])
        ])
        
        logger.info(f"✅ {chain_name}チェーン実行完了")
        return ai_msg["answer"]
        
    except Exception as e:
        logger.error(f"❌ {chain_name}チェーン実行エラー: {e}")
        # エラー時も安全に処理を継続
        return f"申し訳ございませんが、{chain_name}に関する情報の取得でエラーが発生しました。"

def build_knowledge_vectorstore():
    """
    スプレッドシートベースのベクトルDB構築（FAISS版）
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info("🧱 スプレッドシートベースのベクトルDB構築開始")
    
    try:
        # スプレッドシートからQ&Aデータを取得
        docs = load_qa_from_google_sheet(ct.GOOGLE_SHEET_URL)
        
        if not docs:
            logger.warning("⚠️ スプレッドシートからドキュメントが取得できませんでした")
            return False
        
        # 埋め込みオブジェクトを作成
        embeddings = OpenAIEmbeddings()
        
        # FAISSベクトルストアを作成
        db = FAISS.from_documents(docs, embeddings)
        
        # ベクトルストアを保存
        base_path = f"{ct.DB_KNOWLEDGE_PATH}_faiss"
        success = save_faiss_index(db, base_path)
        
        if success:
            # メタデータも保存
            save_index_metadata(base_path, docs)
            logger.info(f"✅ ベクトルDB構築完了: {len(docs)} docs → {base_path}")
            return True
        else:
            logger.error("❌ ベクトルDB保存に失敗しました")
            return False
            
    except Exception as e:
        logger.error(f"❌ ベクトルDB構築中にエラー: {type(e).__name__} - {e}")
        logger.error(f"詳細エラー: {traceback.format_exc()}")
        return False

def load_qa_from_google_sheet(sheet_url: str) -> List[Document]:
    """
    GoogleシートからQ&Aデータを読み込み（修正版）
    
    Args:
        sheet_url: GoogleシートのURL
        
    Returns:
        Documentオブジェクトのリスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info("🔍 スプレッドシート読み込み開始")
    
    try:
        # Google Sheets API認証
        scope = [
            "https://spreadsheets.google.com/feeds", 
            "https://www.googleapis.com/auth/drive"
        ]
        
        # 認証ファイルの存在確認
        auth_file_path = 'secrets/service_account.json'
        if not os.path.exists(auth_file_path):
            logger.error(f"❌ 認証ファイルが見つかりません: {auth_file_path}")
            return []
        
        creds = ServiceAccountCredentials.from_json_keyfile_name(auth_file_path, scope)
        client = gspread.authorize(creds)
        
        # シートを開いてデータを取得
        sheet = client.open_by_url(sheet_url).sheet1
        rows = sheet.get_all_records()
        
        logger.info(f"📊 スプレッドシートから {len(rows)} 行のデータを取得")
        
    except Exception as e:
        logger.error(f"❌ スプレッドシート読み込み失敗: {type(e).__name__} - {e}")
        logger.error(f"詳細エラー: {traceback.format_exc()}")
        return []

    # ドキュメントオブジェクトの作成
    docs = []
    for i, row in enumerate(rows):
        try:
            # 各カラムからデータを取得（柔軟にフィールド名を処理）
            q = row.get("質問", row.get("question", ""))
            a = row.get("回答", row.get("answer", ""))
            src = row.get("根拠資料", row.get("source", ""))
            cat = row.get("対応カテゴリ", row.get("category", ""))
            
            # 必須フィールドのチェック
            if not q or not a:
                logger.warning(f"⚠️ {i+1}行目: 質問または回答が空です (Q='{q}', A='{a}')")
                continue
            
            # メタデータの作成
            meta = {
                "file_name": "GoogleSheet",
                "category": cat if cat else "一般",
                "source": src if src else "社内Q&A",
                "top_keywords": q,  # 質問をキーワードとして使用
                "row_number": i + 1,
                "sheet_url": sheet_url
            }
            
            # コンテンツの作成
            content = f"Q: {q}\nA: {a}"
            if src:
                content += f"\n根拠: {src}"
            
            # Documentオブジェクトを作成
            doc = Document(page_content=content, metadata=meta)
            docs.append(doc)
            
        except Exception as e:
            logger.warning(f"⚠️ {i+1}行目のパース失敗: {type(e).__name__} - {e}")
            continue
    
    logger.info(f"✅ {len(docs)}件の有効なDocumentを作成")
    return docs

def search_knowledge(query: str, top_k: int = 3, score_threshold: float = 0.3):
    """
    スプレッドシートベースの検索（FAISS版）
    
    Args:
        query: 検索クエリ
        top_k: 取得する最大件数
        score_threshold: スコアの閾値
        
    Returns:
        検索結果のDocumentリスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🔎 スプレッドシートベースの検索開始：'{query}'")
    
    try:
        # ナレッジベクトルストアの存在確認
        if not hasattr(st.session_state, 'knowledge_doc_chain'):
            logger.error("❌ knowledge_doc_chainが初期化されていません")
            return []
        
        # 検索実行
        retriever = st.session_state.knowledge_doc_chain.retriever
        results = retriever.get_relevant_documents(query)
        
        if not results:
            logger.info("🔸 検索結果：0件")
            return []
        
        # スコアフィルタリング（FAISSの場合はスコアが取得できない場合があるため柔軟に処理）
        filtered = []
        for doc in results[:top_k]:  # まず上位top_k件に絞る
            score = doc.metadata.get("score", 1.0)  # スコアがない場合は1.0とする
            if score >= score_threshold:
                filtered.append(doc)
        
        # フィルター結果が空の場合は上位結果をそのまま返す
        if not filtered and results:
            filtered = results[:top_k]
            logger.info(f"🔸 スコアフィルター後が空のため、上位{len(filtered)}件を返します")
        
        logger.info(f"🔸 最終検索結果: {len(filtered)}件")
        return filtered
        
    except Exception as e:
        logger.error(f"❌ スプレッドシート検索中にエラー発生: {type(e).__name__} - {e}")
        logger.error(f"詳細エラー: {traceback.format_exc()}")
        return []

def safe_get_secret(key):
    """
    環境変数またはStreamlit Secretsから安全に値を取得
    
    Args:
        key: 環境変数名
        
    Returns:
        値が存在する場合は値、存在しない場合はNone
    """
        
    # まず環境変数から取得を試行
    env_value = os.getenv(key)
    if env_value:
        return env_value
    
    # 環境変数にない場合、Streamlit Secretsから取得を試行
    try:
        return st.secrets.get(key)
    except Exception:
        # secrets.tomlが存在しない、またはキーが存在しない場合
        return None

def check_env_var_status(key):
    """
    環境変数の設定状況を文字列で返す（デバッグ用）
    
    Args:
        key: 環境変数名
        
    Returns:
        "設定済み" または "未設定"
    """
    return "設定済み" if safe_get_secret(key) else "未設定"

def build_error_message(message):
    """
    エラーメッセージと管理者問い合わせテンプレートの連結

    Args:
        message: 画面上に表示するエラーメッセージ

    Returns:
        エラーメッセージと管理者問い合わせテンプレートの連結テキスト
    """
    return "\n".join([message, ct.COMMON_ERROR_MESSAGE])

def clean_keyword(keyword): 
    """
    キーワードから不正な文字を除去
    """
    import re
    # ゼロ幅文字やMarkdown記号を除去
    cleaned = re.sub(r'[\u200b\u200c\u200d\ufeff●○■□▶◇◆\.]', '', keyword)
    # 前後の空白を除去
    cleaned = cleaned.strip()
    return cleaned

def expand_keywords_with_synonyms(keywords: List[str]) -> List[str]:
    """
    キーワードリストに同義語を追加して拡張
    
    Args:
        keywords: 元のキーワードリスト（例: ["受賞歴"]）
        
    Returns:
        同義語を含む拡張されたキーワードリスト（例: ["受賞歴", "受賞", "表彰", "アワード", "賞", "栄誉"]）
    """
    # 重複を避けるためsetを使用
    expanded = set(keywords)
    
    # 各キーワードについて同義語を探す
    for keyword in keywords:
        # 完全一致の同義語を追加
        for base_word, synonyms in SYNONYM_DICT.items():
            if keyword == base_word:
                # base_wordの同義語をすべて追加
                expanded.update(synonyms)
            elif keyword in synonyms:
                # keywordが同義語リストに含まれている場合
                expanded.add(base_word)
                expanded.update(synonyms)
        
        # 部分マッチの同義語を追加
        for base_word, synonyms in SYNONYM_DICT.items():
            if base_word in keyword or keyword in base_word:
                expanded.add(base_word)
                expanded.update(synonyms)
    
    return list(expanded)

def check_flexible_keyword_match(query_keywords: List[str], doc_keywords: List[str]) -> tuple:
    """
    柔軟なキーワードマッチングを実行
    
    Args:
        query_keywords: クエリから抽出されたキーワード
        doc_keywords: 文書から抽出されたキーワード
        
    Returns:
        (マッチしたかどうか, マッチしたキーワードのリスト, マッチタイプ)
    """
    matched_keywords = []  # マッチした組み合わせを記録
    match_types = []       # マッチの種類を記録
    
    # クエリキーワードを同義語で拡張
    expanded_query_keywords = expand_keywords_with_synonyms(query_keywords)
    
    # 拡張されたキーワードと文書キーワードを比較
    for query_kw in expanded_query_keywords:
        for doc_kw in doc_keywords:
            # 1. 完全一致チェック
            if query_kw == doc_kw:
                matched_keywords.append(f"{query_kw}={doc_kw}")
                match_types.append("完全一致")
                continue
            
            # 2. 部分一致チェック（クエリキーワードが文書キーワードに含まれる）
            if query_kw in doc_kw:
                matched_keywords.append(f"{query_kw}⊆{doc_kw}")
                match_types.append("部分一致")
                continue
            
            # 3. 逆部分一致チェック（文書キーワードがクエリキーワードに含まれる）
            if doc_kw in query_kw:
                matched_keywords.append(f"{doc_kw}⊆{query_kw}")
                match_types.append("逆部分一致")
                continue
    
    # マッチしたかどうかを判定
    is_match = len(matched_keywords) > 0
    return is_match, matched_keywords, match_types

def filter_chunks_by_flexible_keywords(docs, query):
    """
    柔軟なキーワードマッチングを使ってチャンクをフィルタリング
    
    Args:
        docs: 検索で得られたチャンクリスト（Documentオブジェクト）
        query: ユーザー入力
        
    Returns:
        フィルター後のチャンクリスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # ステップ1: クエリから名詞を抽出
        tokenizer_obj = dictionary.Dictionary().create()
        mode = tokenizer.Tokenizer.SplitMode.C
        tokens = tokenizer_obj.tokenize(query, mode)
        query_nouns = [
            t.surface() 
            for t in tokens
            if "名詞" in t.part_of_speech() and len(t.surface()) > 1
        ]
        
        logger.info(f"📝 抽出されたクエリ名詞: {query_nouns}")
        
        # ステップ2: 各文書とマッチングを実行
        filtered_docs = []
        for doc in docs:
            # 文書のキーワードを取得
            top_keywords_str = doc.metadata.get("top_keywords", "")
            if top_keywords_str:
                top_keywords = [kw.strip() for kw in top_keywords_str.split(" / ") if kw.strip()]
                
                # 柔軟なキーワードマッチングを実行
                is_match, matched_keywords, match_types = check_flexible_keyword_match(
                    query_nouns, top_keywords
                )
                
                # マッチした場合は結果リストに追加
                if is_match:
                    filtered_docs.append(doc)
                    logger.info(f"✅ マッチ成功:")
                    logger.info(f"   ファイル: {doc.metadata.get('file_name', '不明')}")
                    logger.info(f"   マッチしたキーワード: {matched_keywords}")
                    logger.info(f"   マッチタイプ: {set(match_types)}")
        
        logger.info(f"📊 フィルター結果: {len(docs)} → {len(filtered_docs)} 件")
        
        # ステップ3: フォールバック処理
        if not filtered_docs:
            logger.info("⚠️ フィルター結果が空のため、元のdocsを返します")
            return docs  # 安全策：フィルター結果が空の場合は元のリストを返す
        
        return filtered_docs
        
    except Exception as e:
        logger.error(f"❌ フィルタリング処理でエラーが発生: {e}")
        return docs  # エラー時も安全策として元のリストを返す

def create_knowledge_rag_chain():
    """
    スプレッドシートベースのRAGチェーンを作成（FAISS版）
    
    Returns:
        RAGチェーンオブジェクト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info("🔗 ナレッジRAGチェーン作成開始")
    
    try:
        # FAISSベクトルストアを読み込み
        embeddings = OpenAIEmbeddings()
        base_path = f"{ct.DB_KNOWLEDGE_PATH}_faiss"
        
        # 既存のインデックスを読み込み
        db = load_faiss_index(base_path, embeddings)
        
        if db is None:
            logger.warning("⚠️ 既存のナレッジベクトルストアが見つかりません。新規作成します")
            success = build_knowledge_vectorstore()
            if not success:
                logger.error("❌ ナレッジベクトルストアの作成に失敗しました")
                return None
            
            # 作成後に再度読み込み
            db = load_faiss_index(base_path, embeddings)
            if db is None:
                logger.error("❌ 作成したベクトルストアの読み込みに失敗しました")
                return None
        
        # Retriever作成
        retriever = db.as_retriever(search_kwargs={"k": ct.TOP_K})
        
        # RAGチェーン作成（簡易版）
        question_answer_template = ct.SYSTEM_PROMPT_INQUIRY
        question_answer_prompt = ChatPromptTemplate.from_messages([
            ("system", question_answer_template),
            ("human", "{input}")
        ])
        
        # チェーン作成
        question_answer_chain = create_stuff_documents_chain(
            st.session_state.llm, 
            question_answer_prompt
        )
        
        rag_chain = create_retrieval_chain(retriever, question_answer_chain)
        
        logger.info("✅ ナレッジRAGチェーン作成完了")
        return rag_chain
        
    except Exception as e:
        logger.error(f"❌ ナレッジRAGチェーン作成エラー: {e}")
        logger.error(f"詳細エラー: {traceback.format_exc()}")
        return None


def create_rag_chain(db_name):
    """
    引数として渡されたDB内を参照するRAGのChainを作成（共通化版）

    Args:
        db_name: RAG化対象のデータを格納するデータベース名
        
    Returns:
        RAGチェーンオブジェクト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🔗 RAGチェーン作成開始: {db_name}")
    
    # 共通のベクトルストア作成処理を使用
    db = create_base_vectorstore(db_name)
    retriever = db.as_retriever(search_kwargs={"k": ct.TOP_K})

    question_generator_template = ct.SYSTEM_PROMPT_CREATE_INDEPENDENT_TEXT
    question_generator_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", question_generator_template),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
        ]
    )
    question_answer_template = ct.SYSTEM_PROMPT_INQUIRY
    question_answer_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", question_answer_template),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
        ]
    )

    history_aware_retriever = create_history_aware_retriever(
        st.session_state.llm, retriever, question_generator_prompt
    )

    question_answer_chain = create_stuff_documents_chain(st.session_state.llm, question_answer_prompt)
    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)

    logger.info(f"✅ RAGチェーン作成完了: {db_name}")
    return rag_chain


def create_retriever(db_name):
    """
    指定されたDBパスに基づいてRetrieverのみを作成（共通化版）

    Args:
        db_name: ベクトルDBの保存先ディレクトリ名（または定義名）

    Returns:
        LangChainのRetrieverオブジェクト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🔍 Retriever作成開始: {db_name}")
    
    # 共通のベクトルストア作成処理を使用
    db = create_base_vectorstore(db_name)

    retriever = db.as_retriever(search_kwargs={"k": ct.TOP_K})
    logger.info(f"✅ Retriever作成完了: {db_name}")
    return retriever

def add_docs(folder_path, docs_all):
    """
    フォルダ内のファイル一覧を取得

    Args:
        folder_path: フォルダのパス
        docs_all: 各ファイルデータを格納するリスト
    """
    print(f"📂 読み込もうとしているフォルダ: {folder_path}")
    print(f"📂 フルパス: {os.path.abspath(folder_path)}")

    files = os.listdir(folder_path)
    for file in files:
        file_path = os.path.join(folder_path, file)
        # ファイルの拡張子を取得
        file_extension = os.path.splitext(file)[1]
        # 想定していたファイル形式の場合のみ読み込む
        if file_extension in ct.SUPPORTED_EXTENSIONS:
            # ファイルの拡張子に合ったdata loaderを使ってデータ読み込み
            loader = ct.SUPPORTED_EXTENSIONS[file_extension](f"{folder_path}/{file}")
        else:
            continue

        docs = loader.load()
        p = Path(file_path)

        for doc in docs:
            content = doc.page_content

            # 📌 基本メタ情報
            doc.metadata["file_name"] = p.name
            doc.metadata["file_stem"] = p.stem
            doc.metadata["file_ext"] = p.suffix
            doc.metadata["file_path"] = str(p)
            doc.metadata["category"] = p.parent.name
            doc.metadata["file_mtime"] = datetime.datetime.fromtimestamp(p.stat().st_mtime).isoformat()
            doc.metadata["file_ctime"] = datetime.datetime.fromtimestamp(p.stat().st_ctime).isoformat()

            # 🧑‍💻 作成者（docxの場合のみ取得可能な場合あり）
            if p.suffix == ".docx":
                try:
                    from docx import Document as DocxDocument
                    core_props = DocxDocument(p).core_properties
                    doc.metadata["file_author"] = core_props.author
                except Exception:
                    doc.metadata["file_author"] = "不明"
            else:
                doc.metadata["file_author"] = "不明"

            # 🧩 セクション見出し
            section_titles = []
            for line in content.splitlines():
                if line.strip().startswith(("■", "●", "○", "【", "▶", "◇", "◆")):
                    section_titles.append(line.strip())
            doc.metadata["section_titles"] = " / ".join(section_titles)
            doc.metadata["first_heading"] = section_titles[0] if section_titles else ""
            doc.metadata["section_count"] = len(section_titles)

            # 🧠 頻出キーワード（名詞のみ）- 修正版
            try:
                # 形態素解析の実行
                tokenizer_obj = dictionary.Dictionary().create()
                mode = tokenizer.Tokenizer.SplitMode.C
                tokens = tokenizer_obj.tokenize(content, mode)
                
                # 名詞のみを抽出（クリーニング付き）
                nouns = []
                raw_nouns_sample = []  # デバッグ用サンプル
                
                for i, t in enumerate(tokens):
                    surface = t.surface()
                    if "名詞" in t.part_of_speech() and len(surface) > 1:
                        # デバッグ用サンプル収集（最初の10個まで）
                        if len(raw_nouns_sample) < 10:
                            raw_nouns_sample.append(surface)
                        
                        # キーワードをクリーニング
                        cleaned_surface = clean_keyword(surface)
                        if cleaned_surface and len(cleaned_surface) > 1:  # 空文字や1文字を除外
                            nouns.append(cleaned_surface)
                
                # デバッグ用出力
                if raw_nouns_sample:
                    print(f"📝 クリーニング前サンプル ({p.name}): {raw_nouns_sample}")
                    print(f"📝 クリーニング後サンプル ({p.name}): {nouns[:10]}")
                
                # 頻出キーワードを取得
                if nouns:
                    word_counts = Counter(nouns)
                    top_keywords = [word for word, count in word_counts.most_common(5)]
                    print(f"🔑 抽出キーワード ({p.name}): {top_keywords}")
                else:
                    top_keywords = []
                    print(f"⚠️ 名詞が抽出されませんでした: {p.name}")
                
                # メタデータに設定
                doc.metadata["top_keywords"] = " / ".join(top_keywords)
                
            except Exception as e:
                print(f"❌ キーワード抽出エラー ({p.name}): {e}")
                import traceback
                print(f"詳細エラー: {traceback.format_exc()}")
                doc.metadata["top_keywords"] = ""

            # ✏️ 文字数・行数など（追加で役立つ）
            doc.metadata["num_chars"] = len(content)
            doc.metadata["num_lines"] = len(content.splitlines())

        docs_all.extend(docs)

def run_company_doc_chain(param):
    """
    会社に関するデータ参照に特化したTool設定用の関数

    Args:
        param: ユーザー入力値

    Returns:
        LLMからの回答
    """
    return run_doc_chain_base("company", param)

def run_service_doc_chain(param):
    """
    サービスに関するデータ参照に特化したTool設定用の関数

    Args:
        param: ユーザー入力値

    Returns:
        LLMからの回答
    """
    return run_doc_chain_base("service", param)

def run_customer_doc_chain(param):
    """
    顧客とのやり取りに関するデータ参照に特化したTool設定用の関数

    Args:
        param: ユーザー入力値
    
    Returns:
        LLMからの回答
    """
    return run_doc_chain_base("customer", param)

def run_manual_doc_chain(param):
    """
    操作マニュアルや手順書に関するデータ参照に特化したTool設定用の関数

    Args:
        param: ユーザー入力値

    Returns:
        LLMからの回答
    """
    return run_doc_chain_base("manual", param)

def run_policy_doc_chain(param):
    """
    利用規約・キャンセル・返品などの制度・ポリシーに関するデータ参照に特化したTool設定用の関数

    Args:
        param: ユーザー入力値

    Returns:
        LLMからの回答
    """
    return run_doc_chain_base("policy", param)


def run_sustainability_doc_chain(param):
    """
    環境・サステナビリティ・エシカル活動に関するデータ参照に特化したTool設定用の関数

    Args:
        param: ユーザー入力値

    Returns:
        LLMからの回答
    """
    return run_doc_chain_base("sustainability", param)

def delete_old_conversation_log(result):
    """
    古い会話履歴の削除

    Args:
        result: LLMからの回答
    """
    # LLMからの回答テキストのトークン数を取得
    response_tokens = len(st.session_state.enc.encode(result))
    # 過去の会話履歴の合計トークン数に加算
    st.session_state.total_tokens += response_tokens

    # トークン数が上限値を下回るまで、順に古い会話履歴を削除
    while st.session_state.total_tokens > ct.MAX_ALLOWED_TOKENS:
        # 最も古い会話履歴を削除
        removed_message = st.session_state.chat_history.pop(1)
        # 最も古い会話履歴のトークン数を取得
        removed_tokens = len(st.session_state.enc.encode(removed_message.content))
        # 過去の会話履歴の合計トークン数から、最も古い会話履歴のトークン数を引く
        st.session_state.total_tokens -= removed_tokens

def notice_slack(chat_message):
    """
    問い合わせ内容のSlackへの通知（最適化版）

    Args:
        chat_message: ユーザーメッセージ

    Returns:
        問い合わせサンクスメッセージ
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info("🚀 Slack通知処理を開始します")

    try:
        # === 遅延初期化チェック（Slack用） ===
        if "agent_executor" not in st.session_state:
            logger.info("🔄 Slack処理のため遅延初期化を実行します")
            try:
                from initialize import initialize_heavy_components
                with st.spinner("システム初期化中..."):
                    initialize_heavy_components()
            except Exception as init_error:
                logger.error(f"❌ 遅延初期化に失敗: {init_error}")
                return "お問い合わせを受け付けましたが、システムエラーが発生しました。直接お電話でお問い合わせください。"

        # === Step 1: 担当者選定 ===
        logger.info("👥 担当者選定を開始")
        target_employees = select_responsible_employees(chat_message)
        
        # === Step 2: SlackID取得と通知対象の決定 ===
        if target_employees:
            # 適切な担当者が見つかった場合
            slack_ids = get_slack_ids(target_employees)
            slack_id_text = create_slack_id_text(slack_ids)
            logger.info(f"📧 通知対象SlackID: {slack_id_text}")
            notification_type = "specific_users"
        else:
            # 適切な担当者が見つからない場合は@channelで全員に通知
            logger.warning("⚠️ 適切な担当者が見つかりませんでした。@channelで全員に通知します")
            slack_id_text = "@channel"
            notification_type = "channel_all"

        # === Step 3: 参考情報取得 ===
        logger.info("📚 参考情報を取得中")
        knowledge_context = get_knowledge_context_for_slack(chat_message)

        # === Step 4: 現在日時取得 ===
        now_datetime = get_datetime()
        user_email = st.session_state.get("user_email", "未入力")

        # === Step 5: Slackメッセージ生成 ===
        logger.info("✍️ Slackメッセージを生成中")
        slack_message = generate_slack_message_with_fallback(
            slack_id_text, chat_message, knowledge_context, 
            now_datetime, user_email, notification_type
        )

        # === Step 6: Slack送信 ===
        logger.info("📤 Slackにメッセージを送信中")
        success = send_to_slack_channel(slack_message, "customer-contact2")
        
        if success:
            logger.info("✅ Slack通知が正常に完了しました")
            return ct.CONTACT_THANKS_MESSAGE
        else:
            logger.error("❌ Slack通知に失敗しました")
            return "お問い合わせを受け付けましたが、システムエラーが発生しました。直接お電話でお問い合わせください。"

    except Exception as e:
        logger.error(f"❌ Slack通知処理でエラー発生: {e}")
        logger.error(f"詳細エラー: {traceback.format_exc()}")
        return "お問い合わせを受け付けましたが、システムエラーが発生しました。直接お電話でお問い合わせください。"


def select_responsible_employees(chat_message):
    """
    問い合わせ内容に基づいて担当者を選定

    Args:
        chat_message: ユーザーメッセージ

    Returns:
        選定された担当者リスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # 従業員情報と履歴を読み込み
        loader = CSVLoader(ct.EMPLOYEE_FILE_PATH, encoding=ct.CSV_ENCODING)
        docs = loader.load()
        loader = CSVLoader(ct.INQUIRY_HISTORY_FILE_PATH, encoding=ct.CSV_ENCODING)
        docs_history = loader.load()

        # データの正規化
        for doc in docs:
            doc.page_content = adjust_string(doc.page_content)
            for key in doc.metadata:
                doc.metadata[key] = adjust_string(doc.metadata[key])

        for doc in docs_history:
            doc.page_content = adjust_string(doc.page_content)
            for key in doc.metadata:
                doc.metadata[key] = adjust_string(doc.metadata[key])

        # 参照データの整形
        docs_all = adjust_reference_data(docs, docs_history)
        
        # Retrieverの作成（Faiss版）
        docs_all_page_contents = [doc.page_content for doc in docs_all]
        embeddings = OpenAIEmbeddings()
        db = FAISS.from_documents(docs_all, embeddings)
        retriever = db.as_retriever(search_kwargs={"k": ct.TOP_K})
        
        bm25_retriever = BM25Retriever.from_texts(
            docs_all_page_contents,
            preprocess_func=preprocess_func,
            k=ct.TOP_K
        )
        
        ensemble_retriever = EnsembleRetriever(
            retrievers=[bm25_retriever, retriever],
            weights=ct.RETRIEVER_WEIGHTS
        )

        # 関連性の高い従業員情報を取得
        employees = ensemble_retriever.invoke(chat_message)
        context = get_context(employees)

        # 担当者ID選定のためのプロンプト実行
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", ct.SYSTEM_PROMPT_EMPLOYEE_SELECTION)
        ])
        
        output_parser = CommaSeparatedListOutputParser()
        format_instruction = output_parser.get_format_instructions()

        messages = prompt_template.format_prompt(
            employee_context=context, 
            query=chat_message, 
            format_instruction=format_instruction
        ).to_messages()
        
        employee_id_response = st.session_state.llm(messages)
        employee_ids = output_parser.parse(employee_id_response.content)

        # 選定された担当者情報を取得
        target_employees = get_target_employees(employees, employee_ids)
        
        logger.info(f"👥 選定された担当者数: {len(target_employees)}")
        return target_employees

    except Exception as e:
        logger.error(f"❌ 担当者選定でエラー: {e}")
        return []

def generate_slack_message_with_fallback(slack_id_text, query, knowledge_context, now_datetime, user_email, notification_type):
    """
    Slack用のメッセージを生成（@channel対応版）

    Args:
        slack_id_text: メンション対象のSlackID文字列 または "@channel"
        query: 問い合わせ内容
        knowledge_context: 参考情報
        now_datetime: 現在日時
        user_email: ユーザーメールアドレス
        notification_type: "specific_users" または "channel_all"

    Returns:
        生成されたSlackメッセージ
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # 通知タイプに応じてプロンプトを調整
        if notification_type == "channel_all":
            # @channelの場合のプロンプト
            template = ct.SYSTEM_PROMPT_NOTICE_SLACK_CHANNEL
            
            prompt = PromptTemplate(
                input_variables=["query", "knowledge_context", "now_datetime", "user_email"],
                template=template,
            )
            
            prompt_message = prompt.format(
                query=query,
                knowledge_context=knowledge_context,
                now_datetime=now_datetime,
                user_email=user_email,
                GOOGLE_SHEET_URL=ct.GOOGLE_SHEET_URL,
                WEB_URL=ct.WEB_URL
            )
        else:
            # 特定ユーザーの場合は既存のプロンプト
            prompt = PromptTemplate(
                input_variables=["slack_id_text", "query", "knowledge_context", "now_datetime", "user_email"],
                template=ct.SYSTEM_PROMPT_NOTICE_SLACK,
            )
            
            prompt_message = prompt.format(
                slack_id_text=slack_id_text,
                query=query,
                knowledge_context=knowledge_context,
                now_datetime=now_datetime,
                user_email=user_email,
                GOOGLE_SHEET_URL=ct.GOOGLE_SHEET_URL,
                WEB_URL=ct.WEB_URL
            )

        # LLMでメッセージ生成
        response = st.session_state.llm.invoke([{"role": "user", "content": prompt_message}])
        generated_message = response.content if hasattr(response, 'content') else str(response)
        
        # @channelの場合は先頭に追加
        if notification_type == "channel_all":
            generated_message = f"@channel\n\n{generated_message}"
        
        logger.info("✅ Slackメッセージ生成完了")
        return generated_message

    except Exception as e:
        logger.error(f"❌ メッセージ生成エラー: {e}")
        # フォールバック用の簡単なメッセージ
        if notification_type == "channel_all":
            fallback_message = f"""
@channel

【緊急】適切な担当者が特定できない新しいお問い合わせが届きました。

【問い合わせ内容】
{query}

【問い合わせ者メールアドレス】
{user_email}

【日時】
{now_datetime}

どなたか対応可能な方は、このメッセージにリアクションをお願いします。
            """.strip()
        else:
            fallback_message = f"""
新しいお問い合わせが届きました。

【問い合わせ内容】
{query}

【問い合わせ者メールアドレス】
{user_email}

【日時】
{now_datetime}

担当者の皆様、対応をお願いいたします。
            """.strip()
        
        return fallback_message


def send_to_slack_channel(message, channel_name):
    """
    Slackチャンネルにメッセージを送信

    Args:
        message: 送信するメッセージ
        channel_name: 送信先チャンネル名

    Returns:
        bool: 送信成功時True、失敗時False
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # Slack Bot Tokenを取得
        bot_token = safe_get_secret("SLACK_BOT_TOKEN")
        if not bot_token:
            logger.error("❌ SLACK_BOT_TOKENが設定されていません")
            return False

        # Slack WebClient初期化
        client = WebClient(token=bot_token)
        
        # チャンネルにメッセージ送信
        response = client.chat_postMessage(
            channel=f"#{channel_name}",
            text=message,
            username="問い合わせボット",
            icon_emoji=":robot_face:"
        )
        
        if response["ok"]:
            logger.info(f"✅ Slack送信成功: チャンネル #{channel_name}")
            return True
        else:
            logger.error(f"❌ Slack送信失敗: {response.get('error', '不明なエラー')}")
            return False

    except SlackApiError as e:
        logger.error(f"❌ Slack API エラー: {e.response['error']}")
        return False
    except Exception as e:
        logger.error(f"❌ Slack送信でエラー: {e}")
        return False

def get_knowledge_context_for_slack(chat_message):
    """
    Slack通知用の参考情報を取得（Google Sheets + Web検索）

    Args:
        chat_message: ユーザーメッセージ

    Returns:
        参考情報のテキスト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    knowledge_context = ""

    try:
        # === Google Sheets からQ&A取得 ===
        logger.info("📊 Google Sheetsから情報取得中")
        try:
            scope = [
                "https://spreadsheets.google.com/feeds",
                "https://www.googleapis.com/auth/drive"
            ]
            creds = ServiceAccountCredentials.from_json_keyfile_name(
                'secrets/service_account.json', scope
            )
            client = gspread.authorize(creds)
            sheet = client.open_by_url(ct.GOOGLE_SHEET_URL).sheet1
            rows = sheet.get_all_records()

            sheets_context = "【Google Sheetsから取得した社内Q&A】\n"
            for i, row in enumerate(rows[:10], 1):  # 最初の10件まで
                q = row.get("質問", "")
                a = row.get("回答", "")
                source = row.get("根拠資料", "")
                if q and a:
                    sheets_context += f"{i}. Q: {q}\n   A: {a}\n"
                    if source:
                        sheets_context += f"   根拠: {source}\n"
                    sheets_context += "\n"

            knowledge_context += sheets_context + "\n" + "="*50 + "\n"
            logger.info(f"✅ Google Sheets情報取得完了: {len(rows)}件")

        except Exception as e:
            logger.warning(f"⚠️ Google Sheets取得エラー: {e}")
            knowledge_context += "【Google Sheets情報】取得に失敗しました。\n\n"

        # === Web検索（pip-maker.com）===
        logger.info("🌐 Web検索を実行中")
        try:
            search_wrapper = GoogleSearchAPIWrapper()
            search_query = f"site:pip-maker.com {chat_message}"
            web_results = search_wrapper.run(search_query)
            
            web_context = "【pip-maker.comからの検索結果】\n"
            web_context += web_results[:1000] + "...\n\n"  # 最初の1000文字まで
            
            knowledge_context += web_context
            logger.info("✅ Web検索完了")

        except Exception as e:
            logger.warning(f"⚠️ Web検索エラー: {e}")
            knowledge_context += "【Web検索情報】取得に失敗しました。\n\n"

        return knowledge_context

    except Exception as e:
        logger.error(f"❌ 参考情報取得でエラー: {e}")
        return "参考情報の取得に失敗しました。"


def adjust_reference_data(docs, docs_history):
    """
    Slack通知用の参照先データの整形

    Args:
        docs: 従業員情報ファイルの読み込みデータ
        docs_history: 問い合わせ対応履歴ファイルの読み込みデータ

    Returns:
        従業員情報と問い合わせ対応履歴の結合テキスト
    """

    docs_all = []
    for row in docs:
        # 従業員IDの取得
        row_lines = row.page_content.split("\n")
        row_dict = {item.split(": ")[0]: item.split(": ")[1] for item in row_lines}
        employee_id = row_dict["従業員ID"]

        doc = ""

        # 取得した従業員IDに紐づく問い合わせ対応履歴を取得
        same_employee_inquiries = []
        for row_history in docs_history:
            row_history_lines = row_history.page_content.split("\n")
            row_history_dict = {item.split(": ")[0]: item.split(": ")[1] for item in row_history_lines}
            if row_history_dict["従業員ID"] == employee_id:
                same_employee_inquiries.append(row_history_dict)

        if same_employee_inquiries:
            doc += "【従業員情報】\n"
            row_data = "\n".join(row_lines)
            doc += row_data + "\n=================================\n"
            doc += "【この従業員の問い合わせ対応履歴】\n"
            for inquiry_dict in same_employee_inquiries:
                for key, value in inquiry_dict.items():
                    doc += f"{key}: {value}\n"
                doc += "---------------\n"

            new_doc = Document(page_content=doc, metadata={})
        else:
            new_doc = Document(page_content=row.page_content, metadata={})

        docs_all.append(new_doc)
    
    return docs_all

def get_target_employees(employees, employee_ids):
    """
    問い合わせ内容と関連性が高い従業員情報一覧の取得

    Args:
        employees: 問い合わせ内容と関連性が高い従業員情報一覧
        employee_ids: 問い合わせ内容と関連性が「特に」高い従業員のID一覧

    Returns:
        問い合わせ内容と関連性が「特に」高い従業員情報一覧
    """

    target_employees = []
    duplicate_check = []
    target_text = "従業員ID"
    for employee in employees:
        # 従業員IDの取得
        num = employee.page_content.find(target_text)
        employee_id = employee.page_content[num+len(target_text)+2:].split("\n")[0]
        # 問い合わせ内容と関連性が高い従業員情報を、IDで照合して取得（重複除去）
        if employee_id in employee_ids:
            if employee_id in duplicate_check:
                continue
            duplicate_check.append(employee_id)
            target_employees.append(employee)
    
    return target_employees

def get_slack_ids(target_employees):
    """
    SlackIDの一覧を取得

    Args:
        target_employees: 問い合わせ内容と関連性が高い従業員情報一覧

    Returns:
        SlackIDの一覧
    """

    target_text = "SlackID"
    slack_ids = []
    for employee in target_employees:
        num = employee.page_content.find(target_text)
        slack_id = employee.page_content[num+len(target_text)+2:].split("\n")[0]
        slack_ids.append(slack_id)
    
    return slack_ids

def create_slack_id_text(slack_ids):
    """
    SlackIDの一覧を取得

    Args:
        slack_ids: SlackIDの一覧

    Returns:
        SlackIDを「と」で繋いだテキスト
    """
    slack_id_text = ""
    for i, id in enumerate(slack_ids):
        slack_id_text += f"「{id}」"
        # 最後のSlackID以外、連結後に「と」を追加
        if not i == len(slack_ids)-1:
            slack_id_text += "と"
    
    return slack_id_text

def get_context(docs):
    """
    プロンプトに埋め込むための従業員情報テキストの生成
    Args:
        docs: 従業員情報の一覧

    Returns:
        生成した従業員情報テキスト
    """

    context = ""
    for i, doc in enumerate(docs, start=1):
        context += "===========================================================\n"
        context += f"{i}人目の従業員情報\n"
        context += "===========================================================\n"
        context += doc.page_content + "\n\n"

    return context

def get_datetime():
    """
    現在日時を取得

    Returns:
        現在日時
    """

    dt_now = datetime.datetime.now()
    now_datetime = dt_now.strftime('%Y年%m月%d日 %H:%M:%S')

    return now_datetime

def preprocess_func(text):
    """
    形態素解析による日本語の単語分割
    Args:
        text: 単語分割対象のテキスト

    Returns:
        単語分割を実施後のテキスト
    """

    tokenizer_obj = dictionary.Dictionary(dict="full").create()
    mode = tokenizer.Tokenizer.SplitMode.A
    tokens = tokenizer_obj.tokenize(text ,mode)
    words = [token.surface() for token in tokens]
    words = list(set(words))

    return words

def adjust_string(s):
    """
    Windows環境でRAGが正常動作するよう調整
    
    Args:
        s: 調整を行う文字列
    
    Returns:
        調整を行った文字列
    """
    # 調整対象は文字列のみ
    if type(s) is not str:
        return s

    # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
    if sys.platform.startswith("win"):
        s = unicodedata.normalize('NFC', s)
        s = s.encode("cp932", "ignore").decode("cp932")
        return s
    
    # OSがWindows以外の場合はそのまま返す
    return s

def filter_chunks_by_top_keywords(docs, query):
    """
    top_keywords を使ってチャンクをフィルタリング

    Args:
        docs: 検索で得られたチャンクリスト（Documentオブジェクト）
        query: ユーザー入力

    Returns:
        フィルター後のチャンクリスト（条件に一致しない場合は元のdocsを返す）
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    try:
        # クエリから名詞を抽出
        tokenizer_obj = dictionary.Dictionary().create()
        mode = tokenizer.Tokenizer.SplitMode.C
        tokens = tokenizer_obj.tokenize(query, mode)
        query_nouns = set([
            t.surface() 
            for t in tokens
            if "名詞" in t.part_of_speech() and len(t.surface()) > 1
        ])
        
        # デバッグログ出力
        logger.info(f"フィルター対象クエリ名詞: {query_nouns}")
        logger.info(f"フィルター前チャンク件数: {len(docs)}")
        
        # フィルタリング実行
        filtered_docs = []
        for doc in docs:
            top_keywords_str = doc.metadata.get("top_keywords", "")
            if top_keywords_str:
                # top_keywordsが " / " で区切られている場合の処理
                top_keywords = [kw.strip() for kw in top_keywords_str.split(" / ") if kw.strip()]
                
                # クエリの名詞とキーワードの一致チェック
                if any(kw in query_nouns for kw in top_keywords if kw):
                    filtered_docs.append(doc)
                    logger.info(f"マッチしたキーワード: {[kw for kw in top_keywords if kw in query_nouns]}")
        
        logger.info(f"フィルター後チャンク件数: {len(filtered_docs)}")
        
        # フィルター結果が空の場合は元のdocsを返す（fallback）
        if not filtered_docs:
            logger.info("フィルター結果が空のため、元のdocsを返します")
            return docs
        
        return filtered_docs
        
    except Exception as e:
        logger.error(f"フィルタリング処理でエラーが発生: {e}")
        return docs  # エラー時は元のdocsを返す

def execute_agent_or_chain(chat_message):
    """
    AIエージェントもしくはAIエージェントなしのRAGのChainを実行（コールバックエラー修正版）

    Args:
        chat_message: ユーザーメッセージ

    Returns:
        LLMからの回答
    """
    logger = logging.getLogger(ct.LOGGER_NAME)

    # === 遅延初期化チェック（既存のまま） ===
    if "agent_executor" not in st.session_state:
        logger.info("🔄 agent_executorが未初期化のため、遅延初期化を実行します")
        
        init_placeholder = st.empty()
        with init_placeholder:
            st.info("🔄 初回アクセスのため、システムを初期化しています... （30-60秒程度お待ちください）")
        
        try:
            from initialize import initialize_heavy_components
            initialize_heavy_components()
            logger.info("✅ 遅延初期化が完了しました")
            
            with init_placeholder:
                st.success("✅ 初期化完了！回答を生成しています...")
                
        except Exception as init_error:
            logger.error(f"❌ 遅延初期化に失敗: {init_error}")
            with init_placeholder:
                st.error("❌ 初期化に失敗しました")
            return "申し訳ございませんが、システムの初期化に失敗しました。ページを再読み込みしてください。"
        finally:
            import time
            time.sleep(1)
            init_placeholder.empty()

    # === 実行モードの記録 ===
    logger.info(f"🎯 実行モード: {st.session_state.agent_mode}")
    logger.info(f"📝 入力メッセージ: {chat_message}")

    # AIエージェント機能を利用する場合
    if st.session_state.agent_mode == ct.AI_AGENT_MODE_ON:
        logger.info("🤖 AIエージェントモードで実行")
        
        try:
            # === 修正: StreamlitCallbackHandlerの安全な使用 ===
            # コンテナを事前に作成して、コールバック用の場所を確保
            callback_container = st.container()
            
            # エラーハンドリング付きでCallbackHandlerを作成
            try:
                st_callback = StreamlitCallbackHandler(
                    parent_container=callback_container,
                    max_thought_containers=4,  # 思考過程の表示数を制限
                    expand_new_thoughts=True,
                    collapse_completed_thoughts=True
                )
                
                # Agent実行
                result = st.session_state.agent_executor.invoke(
                    {"input": chat_message}, 
                    {"callbacks": [st_callback]}
                )
                response = result["output"]
                
            except Exception as callback_error:
                logger.warning(f"⚠️ StreamlitCallbackHandlerでエラー: {callback_error}")
                logger.info("🔄 コールバックなしでAgent実行にフォールバック")
                
                # コールバックなしで再実行
                result = st.session_state.agent_executor.invoke(
                    {"input": chat_message},
                    {"callbacks": []}  # 空のコールバック
                )
                response = result["output"]
                
                # 手動で思考過程を表示
                with callback_container:
                    st.info("🤖 AIエージェントが複数のツールを使用して回答を生成しました")
                
        except Exception as agent_error:
            logger.error(f"❌ Agent実行でエラー: {agent_error}")
            response = "申し訳ございませんが、AIエージェント処理でエラーが発生しました。通常モードで再試行してください。"
            
    else:
        # === 通常RAGモード（既存のまま） ===
        logger.info("🔍 通常RAGモードで実行 - 柔軟キーワードマッチング適用")
        
        try:
            retriever = create_retriever(ct.DB_ALL_PATH)
            original_docs = retriever.get_relevant_documents(chat_message)
            logger.info(f"📚 通常検索結果: {len(original_docs)}件")

            logger.info("🧠 柔軟なキーワードマッチングを開始")
            filtered_docs = filter_chunks_by_flexible_keywords(original_docs, chat_message)
            logger.info(f"✅ フィルター後: {len(filtered_docs)}件")

            if filtered_docs:
                logger.info("📖 フィルター後文書でRAG実行")
                context = "\n\n".join([doc.page_content for doc in filtered_docs[:ct.TOP_K]])
                
                question_answer_template = ct.SYSTEM_PROMPT_INQUIRY
                messages = [
                    {"role": "system", "content": question_answer_template.format(context=context)},
                    {"role": "user", "content": chat_message}
                ]
                
                response_obj = st.session_state.llm.invoke(messages)
                response = response_obj.content if hasattr(response_obj, 'content') else str(response_obj)
                logger.info("✅ カスタムRAG処理完了")
            else:
                logger.info("⚠️ フィルター結果が空 - 通常RAGチェーンにフォールバック")
                if "rag_chain" not in st.session_state:
                    logger.warning("⚠️ rag_chainも未初期化です")
                    return "申し訳ございませんが、システムが完全に初期化されていません。ページを再読み込みしてください。"
                
                result = st.session_state.rag_chain.invoke({
                    "input": chat_message,
                    "chat_history": st.session_state.chat_history
                })
                response = result["answer"]

            st.session_state.chat_history.extend([
                HumanMessage(content=chat_message),
                AIMessage(content=response)
            ])
            
        except Exception as e:
            logger.error(f"❌ RAG処理でエラー発生: {e}")
            import traceback
            logger.error(f"詳細エラー: {traceback.format_exc()}")
            
            try:
                if "rag_chain" in st.session_state:
                    result = st.session_state.rag_chain.invoke({
                        "input": chat_message,
                        "chat_history": st.session_state.chat_history
                    })
                    response = result["answer"]
                    logger.info("🔄 フォールバック処理完了")
                else:
                    logger.error("❌ rag_chainも存在しないため、フォールバックも不可能")
                    response = "申し訳ございませんが、現在システムに問題が発生しています。"
            except Exception as e2:
                logger.error(f"❌ フォールバック処理もエラー: {e2}")
                response = "申し訳ございませんが、現在システムに問題が発生しています。"

    # フラグ設定
    if response != ct.NO_DOC_MATCH_MESSAGE:
        st.session_state.answer_flg = True

    logger.info(f"📤 最終回答: {response[:100]}...")
    return response

def test_keyword_filter():
    """
    キーワードフィルターのテスト関数
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    test_queries = [
        "SNS投稿に関する特典はありますか？",
        "海外配送は対応していますか？", 
        "地域貢献活動はありますか？",
        "受賞歴を教えてください",
        "株主優待制度について教えて",
        "環境への取り組みを知りたい",
        "サブスクリプションプランの料金は？"
    ]
    
    retriever = create_retriever(ct.DB_ALL_PATH)
    
    for query in test_queries:
        print(f"\n{'='*100}")
        print(f"🧪 テストクエリ: {query}")
        print('='*100)
        
        # 通常検索
        original_docs = retriever.get_relevant_documents(query)
        
        # フィルター適用
        filtered_docs = filter_chunks_by_top_keywords(original_docs, query)
        
        print(f"📊 結果:")
        print(f"   - 通常検索: {len(original_docs)}件")
        print(f"   - フィルター後: {len(filtered_docs)}件")
        print(f"   - フィルター効果: {(1-len(filtered_docs)/len(original_docs))*100:.1f}%削減")

def test_flexible_keyword_filter():
    """
    柔軟なキーワードフィルターのテスト関数
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    # 問題のあったクエリを含むテストケース
    test_queries = [
        "受賞歴を教えてください",  # メインの問題クエリ
        "SNS投稿に関する特典はありますか？",
        "海外配送は対応していますか？", 
        "地域貢献活動はありますか？",
        "株主優待制度について教えて",
        "環境への取り組みを知りたい",
        "会社の実績を教えて",
        "アワードを受賞していますか？"
    ]
    
    retriever = create_retriever(ct.DB_ALL_PATH)
    
    for query in test_queries:
        print(f"\n{'='*100}")
        print(f"🧪 テストクエリ: {query}")
        print('='*100)
        
        # 通常検索
        original_docs = retriever.get_relevant_documents(query)
        
        # 従来のフィルター
        old_filtered_docs = filter_chunks_by_top_keywords(original_docs, query)
        
        # 新しい柔軟なフィルター
        new_filtered_docs = filter_chunks_by_flexible_keywords(original_docs, query)
        
        print(f"📊 結果比較:")
        print(f"   - 通常検索: {len(original_docs)}件")
        print(f"   - 従来フィルター: {len(old_filtered_docs)}件")
        print(f"   - 柔軟フィルター: {len(new_filtered_docs)}件")
        print(f"   - 改善効果: {len(new_filtered_docs) - len(old_filtered_docs):+d}件")
        
@st.cache_resource
def create_cached_retriever(db_path):
    """
    キャッシュ機能付きのRetriever作成
    
    Args:
        db_path: データベースパス
        
    Returns:
        キャッシュされたRetrieverオブジェクト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🗃️ キャッシュ付きRetriever作成: {db_path}")
    
    return create_retriever(db_path)

@st.cache_resource
def create_cached_rag_chain(db_path):
    """
    キャッシュ機能付きのRAGチェーン作成
    
    Args:
        db_path: データベースパス
        
    Returns:
        キャッシュされたRAGチェーンオブジェクト
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info(f"🔗 キャッシュ付きRAGチェーン作成: {db_path}")
    
    return create_rag_chain(db_path)

def run_lightweight_debug():
    """
    軽量化されたデバッグ処理（起動時間短縮用）
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    logger.info("🔧 軽量デバッグモード実行中...")
    
    try:
        # 最小限のテストのみ実行
        test_query = "受賞歴を教えてください"
        
        # 形態素解析のテスト
        from sudachipy import tokenizer, dictionary
        tokenizer_obj = dictionary.Dictionary().create()
        mode = tokenizer.Tokenizer.SplitMode.C
        tokens = tokenizer_obj.tokenize(test_query, mode)
        nouns = [t.surface() for t in tokens if "名詞" in t.part_of_speech()]
        
        logger.info(f"🧪 軽量デバッグ完了: 抽出名詞 {nouns}")
        
        # フラグ設定
        st.session_state.retriever_debug_done = True
        st.session_state.flexible_keyword_debug_done = True
        
    except Exception as e:
        logger.warning(f"⚠️ 軽量デバッグでエラー: {e}")